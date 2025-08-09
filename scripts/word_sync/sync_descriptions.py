import openpyxl
import docx.document
import re
from .heading_tree import HeadingTree, build_heading_tree
from typing import List, Iterator
from table_generation import Component
from utils.xls_parsing import parse_components_cached, get_description, set_description, set_component_name, get_xls_from_process_type
from utils.xml import insert_paragraph_after
from rapidfuzz import process, fuzz
from .file_manager import WordFileManager, ExcelFileManager
from os import fspath

def get_descriptions(doc : docx.document.Document) -> Iterator[HeadingTree]:
    root = build_heading_tree(doc)
    yield from root.filter(lambda node : node.heading is not None and node.heading.text == "Description")

def similarity_to_rgb(score):
    t = score / 100

    r = int((1 - t) * 255)
    g = int(t * t * 255)
    b = 0

    return r, g, b

def colored_similarity_text(score):
    r, g, b = similarity_to_rgb(score)
    return f"\033[38;2;{r};{g};{b}m({score:.2f}% similar)\033[0m"

class _WordDescription:
    def __init__(self, node : HeadingTree):
        self.node = node
        self.process_type = self.node.get_parent_heading_absolute(1).text # Top level heading is process type
        self.component_name = self.node.get_parent_heading_relative(1).text # One step above is component name

    def description_paragraph(self):
        if len(self.node.paragraphs) > 0:
            # If there exist a description or a placeholder for it, return the paragraph
            return self.node.paragraphs[0]
        else:
            # Otherwise add a paragraph and return it
            para = insert_paragraph_after(self.node.heading, style="Body Text") #type: ignore
            self.node.paragraphs.append(para)
            return para

    def set_component_name_heading(self, text : str):
        print(f"Changed {self.node.parent.heading.text} to {text}")
        self.node.parent.heading.text = text

class DescriptionSyncer:
    def __init__(self):
        self._process_to_xls_path = {}
        self._xls_managers = {}
        self._word_manager = None

    def sync_descriptions(self, doc_path : str, xls_file_paths : List[str]):
        self._word_manager = WordFileManager(doc_path)

        for desc in get_descriptions(self._word_manager.doc):
            description = _WordDescription(desc)

            xls_path = get_xls_from_process_type(description.process_type, xls_file_paths)
            if xls_path is None:
                print(f"Skipping {description.process_type}")
                continue # Skip iteration if no xls file is chosen
            
            xls_manager = self._parse_excel(xls_path)
            components = parse_components_cached(xls_manager.xls)
            best_matching_component = self._find_best_component_match(description, components)

            if best_matching_component is None:
                print(f"Skipping {description.component_name}")
                continue

            self._set_descriptions(description, best_matching_component, xls_manager.wb)

    def save_files(self):
        if self._word_manager is not None:
            self._word_manager.backup_and_save()
        for xls_manager in self._xls_managers.values():
            xls_manager.backup_and_save()

    def _set_descriptions(self, description : _WordDescription, component : Component, wb : openpyxl.Workbook):
        paragraph = description.node.get_or_insert_paragraph(0)

        word_description = paragraph.text
        excel_description = get_description(wb, component.id)

        if word_description == excel_description:
            return

        if len(word_description) > len(excel_description):
            print(f"Syncing description for {component.id} using Word description. '{excel_description}' -> '{word_description}'")
            set_description(wb, component.id, word_description)
        elif len(excel_description) > len(word_description):
            print(f"Syncing description for {component.id} using Excel description. '{word_description}' -> '{excel_description}'")
            paragraph.text = excel_description

    def _parse_excel(self, xls_path : str) -> ExcelFileManager:
        if xls_path in self._xls_managers:
            return self._xls_managers[xls_path]
        
        xls_manager = ExcelFileManager(xls_path)
        self._xls_managers[xls_path] = xls_manager
        return xls_manager

    def _find_best_component_match(self, description : _WordDescription, components : List[Component]) -> Component | None:
        target = description.component_name

        choices = [comp.name for comp in components] # Extract names of components
        best_match = process.extractOne(
            target,
            choices,
            scorer=fuzz.ratio
        )

        _, similarity, index = best_match
        best_matching_component = components[index]

        if int(similarity) == 100:
            return best_matching_component
        return self._handle_component_mismatch(description, best_matching_component, similarity)

    def _handle_component_mismatch(self, description : _WordDescription, component : Component, similarity) -> Component | None:
        sim_text = colored_similarity_text(similarity)
        print(f"\033[31mFound mismatch \033[0m{sim_text}:")
        print(f"\tIn '{description.process_type}'")
        print(f"\tWord:  '{description.component_name}'")
        print(f"\tExcel: '{component.name}'")

        while True:
            match input("Replace all instances with Word (W), Excel (E), or skip this heading (S)\n-> ").lower():
                case "w":
                    self._set_headings(description, component, description.component_name)
                    return component
                case "e":
                    self._set_headings(description, component, component.name)
                    return component
                case "s" | "":
                    return None
                case _:
                    print("Unkown command, use (W), (E), or (S)")

    def _set_headings(self, description : _WordDescription, component : Component, text : str):
        description.set_component_name_heading(text)

        component_file_path = fspath(component.xls)
        file_manager = self._parse_excel(component_file_path)
        set_component_name(file_manager.wb, component, text)