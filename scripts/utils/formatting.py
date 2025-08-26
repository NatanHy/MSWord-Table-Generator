from typing import TYPE_CHECKING, Any

import docx.document
from docx import Document
from docx.shared import Cm, Pt # Pt import needed so it can be used by the eval when applpying styling
from docx.table import _Cell
from docx.text.paragraph import Paragraph

from config.document_config import TABLE_HEADING_STYLE, LEFT_MARGIN, RIGHT_MARGIN
from utils.xml import insert_multilevel_table_caption, clear_document, insert_paragraph_after
if TYPE_CHECKING:
    from table_generation import FixedTable, Component

def copy_document_styles(path) -> docx.document.Document:
    """
    Copies the formatting and styles from an existing document and returns a blank document.
    """
    template = Document(path)
    clear_document(template)
    return template

def add_table_heading(doc : docx.document.Document, component : 'Component', insert_after=None) -> Paragraph:
    if TABLE_HEADING_STYLE in doc.styles:
        styl = TABLE_HEADING_STYLE
    else:
        styl = None

    if insert_after is None:
        caption = doc.add_paragraph("", style=styl)
    else:
        caption = insert_paragraph_after(insert_after, text="", style=styl)

    table_text = f"Direct dependencies between the process “{component.name}” and the defined {component.system_component} variables."
    insert_multilevel_table_caption(caption, table_text)
    return caption

def format_document(doc : docx.document.Document): 
    sections = doc.sections

    for section in sections:
        section.left_margin = Cm(LEFT_MARGIN)
        section.right_margin = Cm(RIGHT_MARGIN)

def format_table(table : 'FixedTable', format : str):
    _apply_attributes(table, format)

def format_raw_value(val : Any) -> str:
    if val is None:
        return ""

    # Ugly nan-check, but avoids additional conversion
    match str(val):
        case "nan" | "0":
            return "" 
        case res:
            return res
        
def style(cell : _Cell, style : str):
    paragraphs = cell.paragraphs
    for paragraph in paragraphs:
        for run in paragraph.runs:
            _apply_attributes(run.font, style)

def _apply_attributes(font, attr_string):
    pairs = attr_string.split(',')
    for pair in pairs:
        if not pair.strip():
            continue
        try:
            key, value_expr = pair.split('=', 1)
            key = key.strip()

            # Raw eval, if we ever expect an untrusted user to modify table.dsl, this should be changed
            value = eval(value_expr.strip())
            setattr(font, key, value)
        except Exception as e:
            print(f"Error processing '{pair}': {e}")

