from docx.text.paragraph import Paragraph
from docx import Document
import docx.document
from typing import List, Iterator
import re

class HeadingTree:
    def __init__(self, heading : str, paragraphs : List[Paragraph] | None = None):
        self.heading = heading
        if paragraphs is None:
            self.paragraphs = []
        else:
            self.paragraphs = paragraphs
        self.children = []
        self.parent = None

    def add_paragraph(self, paragraph : Paragraph):
        self.paragraphs.append(paragraph)

    def add_child(self, child : 'HeadingTree'):
        child.parent = self
        self.children.append(child)

def get_heading_level(style_name: str) -> int | None:
    """Extract heading level from style name, e.g., 'Heading 2' -> 2."""
    match = re.match(r"Heading (\d+)", style_name)
    return int(match.group(1)) if match else None

def build_heading_tree(doc: docx.document.Document) -> HeadingTree:
    root = HeadingTree("Document Root")
    stack = [(0, root)]  # Each item is a tuple (level, node)

    for para in doc.paragraphs:
        level = get_heading_level(para.style.name) #type: ignore

        if level is not None:
            # Create new node
            new_node = HeadingTree(para.text)

            # Find parent in the stack (last one with lower level)
            while stack and stack[-1][0] >= level:
                stack.pop()

            _, parent_node = stack[-1]
            parent_node.add_child(new_node)
            stack.append((level, new_node))
        else:
            # Add non-heading paragraph to the current top node
            if stack:
                stack[-1][1].add_paragraph(para)

    return root

def filter_descriptions(node: HeadingTree) -> Iterator[HeadingTree]:
    if node.heading == "Description":
        yield node
    for child in node.children:
        yield from filter_descriptions(child)

def get_descriptions(doc : docx.document.Document) -> Iterator[HeadingTree]:
    tree = build_heading_tree(doc)
    yield from filter_descriptions(tree)
