import openpyxl
import docx.document
import re
from .heading_tree import HeadingTree, build_heading_tree
from typing import List, Iterator, Generator
from table_generation import Component
from utils.xls_parsing import parse_components_cached, get_description, set_description, set_component_name, get_xls_from_process_type
from utils.xml import insert_paragraph_after
from rapidfuzz import process, fuzz
from .file_manager import WordFileManager, ExcelFileManager
from os import fspath
from dataclasses import dataclass

def get_descriptions(doc : docx.document.Document) -> Iterator[HeadingTree]:
    root = build_heading_tree(doc)
    yield from root.filter(lambda node : node.heading is not None and node.heading.text == "Description")

@dataclass
class Mismatch:
    mismatch_type : str
    similarity : float
    header : str
    in_word : str
    in_excel : str

class _WordDescription:
    def __init__(self, node : HeadingTree):
        self.node = node
        self.process_type = self.node.get_parent_heading_absolute(1).text #type: ignore Top level heading is process type
        self.component_name = self.node.get_parent_heading_relative(1).text #type: ignore One step above is component name

    def description_paragraph(self):
        if len(self.node.paragraphs) > 0:
            # If there exist a description or a placeholder for it, return the paragraph
            return self.node.paragraphs[0]
        else:
            # Otherwise add a paragraph and return it
            para = insert_paragraph_after(self.node.heading, style="Body Text") #type: ignore
            self.node.paragraphs.append(para)
            return para

    def set_component_name_heading(self, text : str):
        self.node.parent.heading.text = text #type: ignore


class WordExcelSyncer:
    def __init__(self):
        self._process_to_xls_path = {}
        self._xls_managers = {}
        self._word_manager = None

    def sync_files(self, doc_path : str, xls_file_paths : List[str]) -> Generator[Mismatch, str, None]:
        """
        Sync descriptions between a word document and excel file. Also allows syncing of
        mismatched component names in headers. 
        """
        self._word_manager = WordFileManager(doc_path)

        for desc in get_descriptions(self._word_manager.doc):
            description = _WordDescription(desc)

            xls_path = get_xls_from_process_type(description.process_type, xls_file_paths)
            if xls_path is None:
                continue # Skip iteration if no matching xls file is found
            
            xls_manager = self._parse_excel(xls_path)
            components = parse_components_cached(xls_manager.xls)
            best_matching_component = yield from self._find_best_component_match(description, components)

            if best_matching_component is None:
                continue
            
            yield from self._set_descriptions(description, best_matching_component, xls_manager.wb)

    def save_files(self):
        if self._word_manager is not None:
            self._word_manager.backup_and_save()
        for xls_manager in self._xls_managers.values():
            xls_manager.backup_and_save()

    def _set_descriptions(self, description : _WordDescription, component : Component, wb : openpyxl.Workbook) -> Generator[Mismatch, str, None]:
        paragraph = description.node.get_or_insert_paragraph(0)

        word_description = paragraph.text
        excel_description = get_description(wb, component.id)

        similarity = fuzz.ratio(word_description, excel_description)

        while True:
            choice = yield Mismatch(
                "description", 
                similarity, 
                f"{description.process_type.strip()} - {description.component_name.strip()}", 
                word_description, 
                excel_description
                )

            match choice:
                case "w":
                    set_description(wb, component.id, word_description)
                    return
                case "e":
                    paragraph.text = excel_description
                    return
                case "s" | "": 
                    return
                case _:
                    # Unkown command
                    pass

    def _parse_excel(self, xls_path : str) -> ExcelFileManager:
        if xls_path in self._xls_managers:
            return self._xls_managers[xls_path]
        
        xls_manager = ExcelFileManager(xls_path)
        self._xls_managers[xls_path] = xls_manager
        return xls_manager

    def _find_best_component_match(self, description : _WordDescription, components : List[Component]) -> Generator[Mismatch, str, Component | None]:
        target = description.component_name.strip()

        choices = [comp.name for comp in components] # Extract names of components
        best_match = process.extractOne(
            target,
            choices,
            scorer=fuzz.ratio
        )

        _, similarity, index = best_match
        best_matching_component = components[index]

        if int(similarity) == 100:
            return best_matching_component
        
        handled_mismatch = yield from self._handle_component_mismatch(description, best_matching_component, similarity)
        return handled_mismatch

    def _handle_component_mismatch(self, description : _WordDescription, component : Component, similarity) -> Generator[Mismatch, str, Component | None]:
        while True:
            choice = yield Mismatch("component", similarity, description.process_type, description.component_name, component.name)
            match choice:
                case "w":
                    self._set_headings(description, component, description.component_name)
                    return component
                case "e":
                    self._set_headings(description, component, component.name)
                    return component
                case "s" | "":
                    return None
                case _:
                    # Unknown command
                    pass

    def _set_headings(self, description : _WordDescription, component : Component, text : str):
        description.set_component_name_heading(text)

        component_file_path = fspath(component.xls)
        file_manager = self._parse_excel(component_file_path)
        set_component_name(file_manager.wb, component, text)