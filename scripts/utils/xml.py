from typing import cast, Union, Dict, Iterable, Tuple

import docx.document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

BlockItem = Union[Paragraph, Table]

def delete_paragraph(paragraph):
    """
    Deletes the given paragraph from the document.
    """
    p = paragraph._element
    parent = p.getparent()
    parent.remove(p)

def remove_table_after_heading(doc, heading_text):
    """
    Finds a heading by text, then removes the immediate table after it
    (only if it's the very next block).
    """
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == heading_text:
            return _remove_immediate_table_after_paragraph(paragraph)
    return False

def insert_table_after(rows : int, cols : int, insert_after) -> Table:
    # Create the table (in memory, detached)
    tbl = OxmlElement('w:tbl')

    # Add tblPr (optional, for compatibility)
    tblPr = OxmlElement('w:tblPr')
    tbl.append(tblPr)

    # Add table grid (optional)
    tblGrid = OxmlElement('w:tblGrid')
    for _ in range(cols):
        gridCol = OxmlElement('w:gridCol')
        tblGrid.append(gridCol)
    tbl.append(tblGrid)

    # Add rows
    for _ in range(rows):
        tr = OxmlElement('w:tr')
        for _ in range(cols):
            tc = OxmlElement('w:tc')
            p = OxmlElement('w:p')
            tc.append(p)
            tr.append(tc)
        tbl.append(tr)

    # Insert into document
    parent = insert_after._element.getparent()
    index = parent.index(insert_after._element)
    parent.insert(index + 1, tbl)
    tbl = cast(CT_Tbl, tbl)
    return Table(tbl, insert_after._parent)

def insert_paragraph_after(item : BlockItem, text=None, style=None):
    """
    Insert a new paragraph after the given block-level item (Paragraph or Table).
    """
    # Get the correct XML element
    element = item._element

    # Create and insert new paragraph XML
    new_p = OxmlElement("w:p")
    element.addnext(new_p)

    # Wrap in Paragraph with correct parent
    new_para = Paragraph(new_p, item._parent)

    if text:
        new_para.add_run(text)
    if style is not None:
        new_para.style = style

    return new_para

def clear_document(doc):
    body = doc.element.body

    # Save the last section definition if it exists
    sectPr = None
    for el in reversed(body):
        if el.tag == qn('w:sectPr'):
            sectPr = el
            break

    # Remove all children from the body
    for el in list(body):
        body.remove(el)

    # Re-append the final section definition
    if sectPr is not None:
        body.append(sectPr)

def insert_multilevel_table_caption(paragraph, table_text):
    """
    Adds multi-level caption to table, e.g Table 1-1 or Table 1-2, by modifying the xml.
    """
    paragraph.add_run("Table ")

    # Insert STYLEREF 1 \s field for section number
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'STYLEREF 1 \\s'

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')

    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')

    r = paragraph.add_run()
    r._r.append(fldChar1)
    r._r.append(instrText)
    r._r.append(fldChar2)
    r._r.append(fldChar3)

    paragraph.add_run("-")

    # Insert SEQ Table \s 1 for per-section table numbering
    fldChar1_seq = OxmlElement('w:fldChar')
    fldChar1_seq.set(qn('w:fldCharType'), 'begin')

    instrText_seq = OxmlElement('w:instrText')
    instrText_seq.set(qn('xml:space'), 'preserve')
    instrText_seq.text = 'SEQ Table \\* ARABIC \\s 1'

    fldChar2_seq = OxmlElement('w:fldChar')
    fldChar2_seq.set(qn('w:fldCharType'), 'separate')

    fldChar3_seq = OxmlElement('w:fldChar')
    fldChar3_seq.set(qn('w:fldCharType'), 'end')

    r_seq = paragraph.add_run()
    r_seq._r.append(fldChar1_seq)
    r_seq._r.append(instrText_seq)
    r_seq._r.append(fldChar2_seq)
    r_seq._r.append(fldChar3_seq)

    paragraph.add_run(f" {table_text}")

def parse_mappings(doc) -> Dict[str, Dict[str, str]]:
    mappings = {}

    for heading, tbl in get_mapping_tables(doc):
        try:
            mapping = _parse_mapping_table(tbl)
            mappings[heading.text.strip()] = mapping
        except ValueError as e:
            # Ignore error, most likely attempting to parse a non-process type header
            pass

    return mappings

def get_mapping_tables(doc) -> Iterable[Tuple[Paragraph, Table]]:
    for heading, tbl in _get_first_table_per_heading(doc):
        if _is_mapping_table(tbl):
            yield (heading, tbl)

def _is_mapping_table(tbl) -> bool:
    # Check that the header matches the template
    mapping_header = {
        (0, 0): "This report",
        (0, 2): "FSAR FEP catalogue",
        (1, 0): "Section",
        (1, 1): "Process name",
        (1, 2): "FEP ID",
        (1, 3): "FEP Name"
    }
    for indx, s in mapping_header.items():
        if (v := tbl.rows[indx[0]].cells[indx[1]].text.strip()) != s:
            return False

    return True    

def _parse_mapping_table(tbl : Table) -> Dict[str, str]:
    # Map column 1 to column 2
    mapping = {c1.text.strip() : c2.text.strip() for c1, c2 in zip(tbl.column_cells(1), tbl.column_cells(2))}

    # Remove mapping of header
    mapping.pop("This report")
    mapping.pop("Process name")

    return mapping

def _get_first_table_per_heading(doc):
    heading = None
    prev_heading = None

    for block in _iter_block_items(doc):
        if isinstance(block, Paragraph):
            style_name = block.style.name #type: ignore
            if style_name.startswith("Heading"): #type: ignore
                # Extract the level number
                level = int(style_name.split(" ")[1]) #type: ignore
                if level == 1:
                    heading = block
                    continue
                
        elif isinstance(block, Table):
            if heading and heading is not prev_heading:
                prev_heading = heading
                yield (heading, block)

def _iter_block_items(parent):
    """
    Yield paragraphs and tables in document order.
    `parent` can be a Document object or a _Cell object.
    """
    if isinstance(parent, docx.document.Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("Unsupported parent type")

    for child in parent_elm.iterchildren():
        if child.tag.endswith('}p'):  # Paragraph
            yield Paragraph(child, parent)
        elif child.tag.endswith('}tbl'):  # Table
            yield Table(child, parent)

def _remove_immediate_table_after_paragraph(paragraph):
    """
    Removes the table immediately following the given paragraph,
    if and only if the very next block item is a table.
    """
    p_element = paragraph._element
    next_element = p_element.getnext()


    if next_element is not None and next_element.tag.endswith('tbl'):
        parent = next_element.getparent()
        parent.remove(next_element)
        return True  # Removed
    return False  # Nothing removed
