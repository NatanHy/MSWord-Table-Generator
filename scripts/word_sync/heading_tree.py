from typing import List, Iterator, Callable
import re

import docx.document
from docx.text.paragraph import Paragraph
from docx.table import Table

from utils.xml import insert_paragraph_after

def _iter_block_items(parent) -> Iterator[Paragraph | Table]:
    """
    Yield paragraphs and tables in document order.
    Works for docx.Document and docx.table._Cell objects.
    """
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P

    for child in parent.element.body:
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

class HeadingTree:
    def __init__(
            self, 
            heading : Paragraph | None, 
            paragraphs : List[Paragraph] | None = None,
            tables : List[Table] | None = None
            ):
        self.heading = heading

        if paragraphs is None:
            self.paragraphs = []
        else:
            self.paragraphs = paragraphs

        if tables is None:
            self.tables = []
        else:
            self.tables = tables

        self.children : List[HeadingTree] = []
        self.parent : HeadingTree | None = None

    @property
    def level(self) -> int:
        if self.parent is None:
            return 0
        return 1 + self.parent.level

    def get_parent_heading_absolute(self, level : int) -> Paragraph | None:
        steps = self.level - level
        if steps < 0:
            return None
        return self.get_parent_heading_relative(steps)
    
    def get_parent_heading_relative(self, steps : int) -> Paragraph | None:
        if steps == 0:
            return self.heading
        if self.parent is None:
            return None
        return self.parent.get_parent_heading_relative(steps - 1)

    def add_table(self, table : Table):
        self.tables.append(table)

    def add_paragraph(self, paragraph : Paragraph):
        self.paragraphs.append(paragraph)

    def add_child(self, child : 'HeadingTree'):
        child.parent = self
        self.children.append(child)

    def filter(self, key : Callable[['HeadingTree'], bool]) -> Iterator['HeadingTree']:
        if key(self):
            yield self
        for child in self.children:
            yield from child.filter(key)

    def get_or_insert_paragraph(self, indx, style="Body Text"):
        if len(self.paragraphs) > 0:
            # If there exist a paragraph, return it
            return self.paragraphs[indx]
        else:
            # Otherwise add a paragraph and return it
            if self.heading:
                para = insert_paragraph_after(self.heading, style=style)
                self.paragraphs.append(para)
                return para
            else:
                raise RuntimeError("Cannot get paragraph of node because heading is None")

def _get_heading_level(style_name: str) -> int | None:
    """Extract heading level from style name, e.g., 'Heading 2' -> 2."""
    match = re.match(r"Heading (\d+)", style_name)
    return int(match.group(1)) if match else None

def build_heading_tree(doc: docx.document.Document) -> HeadingTree:
    root = HeadingTree(None)
    stack = [(0, root)]  # Each item is a tuple (level, node)

    for block_item in _iter_block_items(doc):
        if isinstance(block_item, Paragraph):
            level = _get_heading_level(block_item.style.name) #type: ignore

            if level is not None:
                # Create new node
                new_node = HeadingTree(block_item)

                # Find parent in the stack (last one with lower level)
                while stack and stack[-1][0] >= level:
                    stack.pop()

                _, parent_node = stack[-1]
                parent_node.add_child(new_node)
                stack.append((level, new_node))
            else:
                # Add non-heading paragraph to the current top node
                if stack:
                    stack[-1][1].add_paragraph(block_item)
        elif isinstance(block_item, Table):
            if stack:
                stack[-1][1].add_table(block_item)

    return root
